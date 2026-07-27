"""
Weekly Equity Report Generator
================================
Monarch Networth Capital
"""

import sys
import os
import warnings
from datetime import datetime, timedelta
from typing import Dict, Any, Optional

import pandas as pd

# Configure console encoding for Windows
if sys.stdout and hasattr(sys.stdout, 'reconfigure'):
    try:
        sys.stdout.reconfigure(encoding='utf-8', errors='replace')
    except Exception:
        pass

warnings.filterwarnings("ignore")

# Local modules
from config import (
    TEMPLATE_PATH, INDEX_NAMES, SECTOR_NAMES, EMA_NAMES, EMA_PERIODS,
    LONG_WINDOW_INDICES, LONG_HISTORY_DAYS, SHORT_HISTORY_DAYS,
    SR_ROUNDING_RULES, SR_FACTORS, COLOR_GREEN, COLOR_RED,
    # Page 2 constants
    CHART_INDICES, HEADING_COLOR, HEADING_FONT_SIZE_HALFPTS,
    CHART_TABLE_WIDTH, CHART_TABLE_COL0_WIDTH, CHART_TABLE_COL1_WIDTH,
    CHART_TABLE_BORDER_COLOR, CHART_TABLE_BORDER_SZ,
    CHART_TABLE_ROW_HEIGHTS, CHART_IMG_CX, CHART_IMG_CY,
    ERT_FONT_SIZE_HALFPTS, ERT_COLOR,
    CONTACT_TABLE_GRID_COLS, CONTACT_TABLE_CELL_WIDTHS,
    CONTACT_TABLE_FONT_SIZE, CONTACT_TABLE_ROW1_HEIGHT,
    RESEARCH_TEAM, DISCLAIMER_TEXT, COMPANY_TEXT, COMPANY_TEXT_FONT_SIZE,
)
from validation import validate_end_sunday, get_friday_dates, find_last_trading_day
from math_utils import (
    round0, round2, calculate_weekly_pct, calculate_ema,
    determine_bias, calculate_support_resistance
)
from data_sources import (
    fetch_with_fallback, get_close_on_date, create_nse_session,
    fetch_nse_snapshot, fetch_fii_dii, fetch_constituents, fetch_global_markets,
    get_weekly_fii_dii, get_fii_dii_for_date,
    parse_bhavcopy_derivatives, fetch_option_chain_live
)


def fetch_all_market_data(start_date, end_date) -> Dict[str, Any]:
    """
    Fetch all market data with proper Friday handling for holidays.

    Returns dict with keys: indices, sectors, ema, sr, data_quality
    """
    print("\n" + "=" * 70)
    print("FETCHING MARKET DATA")
    print("=" * 70)

    # Calculate Friday dates
    prev_friday_target, curr_friday_target = get_friday_dates(start_date, end_date)

    print(f"\nTarget dates:")
    print(f"  Previous week Friday: {prev_friday_target.strftime('%d-%b-%Y')}")
    print(f"  Current week Friday:  {curr_friday_target.strftime('%d-%b-%Y')}")

    # Fetch windows
    fetch_to = end_date + timedelta(days=1)  # yfinance end is exclusive

    # NSE live snapshot (for enrichment only, not for weekly %)
    print("\n[1/3] Fetching NSE live snapshot...")
    nse_session = create_nse_session()
    nse_snap = fetch_nse_snapshot(nse_session)

    # Yahoo Finance historical data
    print("\n[2/3] Downloading Yahoo Finance historical data...")
    yf_cache = {}
    all_symbols = list(dict.fromkeys(INDEX_NAMES + SECTOR_NAMES))

    for disp_name in all_symbols:
        # Use long window for EMA indices
        if disp_name in LONG_WINDOW_INDICES:
            fetch_from = start_date - timedelta(days=LONG_HISTORY_DAYS)
        else:
            fetch_from = start_date - timedelta(days=SHORT_HISTORY_DAYS)

        symbol, df = fetch_with_fallback(disp_name, fetch_from, fetch_to)

        if not df.empty:
            yf_cache[disp_name] = df
            print(f"  [OK] {disp_name:20s}: {len(df):4d} rows [{symbol}]")
        else:
            yf_cache[disp_name] = df  # Empty DataFrame
            print(f"  [FAIL] {disp_name:20s}: No data")

    # Find actual trading days (handle holidays)
    print("\n[3/3] Identifying actual trading days (handling holidays)...")

    # Use Nifty 50 as reference for finding trading days
    nifty_df = yf_cache.get("NIFTY 50")

    if nifty_df is not None and not nifty_df.empty:
        prev_friday_actual = find_last_trading_day(prev_friday_target, nifty_df)
        curr_friday_actual = find_last_trading_day(curr_friday_target, nifty_df)

        # Check if YF is delayed — use NSE live snapshot to confirm the date
        if pd.Timestamp(curr_friday_target) not in nifty_df.index:
            nse_nifty = nse_snap.get("NIFTY 50", {})
            if nse_nifty.get("close") is not None:
                yf_prev_close = get_close_on_date(nifty_df, curr_friday_target - timedelta(days=1))
                if yf_prev_close is not None and abs(nse_nifty["close"] - yf_prev_close) > 1.0:
                    print(f"  [INFO] Yahoo Finance is delayed. NSE live confirms {curr_friday_target.strftime('%d-%b-%Y')} was a trading day.")
                    curr_friday_actual = curr_friday_target

        if prev_friday_actual != prev_friday_target:
            print(f"  [INFO] Previous Friday was holiday — using {prev_friday_actual.strftime('%d-%b-%Y')}")

        if curr_friday_actual != curr_friday_target:
            print(f"  [INFO] Current Friday was holiday or delayed — using {curr_friday_actual.strftime('%d-%b-%Y')}")
    else:
        print("  [WARN] Cannot verify trading days — using target Fridays")
        prev_friday_actual = prev_friday_target
        curr_friday_actual = curr_friday_target

    # ── Helper: resolve the best available current close for a symbol ─────────
    def _resolve_close(disp_name, df):
        """
        Pick the most accurate current close in priority order:
        1. Exact date match in Yahoo Finance data (non-NaN Close)
        2. NSE live snapshot (if YF is missing the date OR Close is NaN)
        3. Latest YF close (if YF data is recent enough — after prev_friday)
        4. None (if YF data is fully stale or all sources are empty)

        Known YF quirk: some indices (FINNIFTY, NIFTYNEXT50) return a row
        for curr_friday with Open/High/Low present but Close = NaN.
        get_close_on_date() then falls back to the last non-NaN close
        which is prev_friday — giving 0% weekly change.  We detect this
        case and explicitly use the NSE live snapshot instead.
        """
        nse_live = nse_snap.get(disp_name, {}).get("close")
        yf_has_exact = False
        yf_curr_nan = False   # True when YF has a row for curr_friday but Close is NaN
        is_stale_week = False

        if df is not None and not df.empty:
            ts = pd.Timestamp(curr_friday_actual)
            valid_df = df.dropna(subset=["Close"])
            yf_has_exact = ts in valid_df.index

            # Detect the YF NaN-close quirk: row exists but Close is NaN
            if not yf_has_exact and ts in df.index:
                yf_curr_nan = True

            # Stale = last valid close is strictly before prev_friday
            # (using < not <= so that data landing exactly on prev_friday
            #  does NOT incorrectly mark the whole fetch as stale)
            is_stale_week = (
                valid_df.empty
                or valid_df.index[-1] < pd.Timestamp(prev_friday_actual)
            )

        yf_curr = get_close_on_date(df, curr_friday_actual)
        yf_prev = get_close_on_date(df, prev_friday_actual)

        if yf_has_exact:
            # YF has a proper (non-NaN) close for curr_friday — use it
            return yf_prev, yf_curr, "YF"
        elif yf_curr_nan or is_stale_week:
            # YF Close is NaN or data is too old — prefer NSE live
            if nse_live is not None:
                return yf_prev, nse_live, "NSE_Fallback"
            else:
                return yf_prev, None, "YF_Stale_Week"
        elif nse_live is not None:
            # YF date is missing entirely — fall back to NSE live
            return yf_prev, nse_live, "NSE_Fallback"
        else:
            # Best-effort: use whatever YF returned
            return yf_prev, yf_curr, "YF_Stale"

    # Build indices data
    print("\n" + "-" * 70)
    print("PROCESSING INDICES")
    print("-" * 70)

    indices = []
    data_quality = {"indices_ok": 0, "indices_total": len(INDEX_NAMES)}

    # resolved_closes stores the corrected curr close for each display name,
    # shared by the EMA and S/R sections below.
    resolved_closes = {}

    for disp_name in INDEX_NAMES:
        df = yf_cache.get(disp_name)
        yf_prev, curr_close, src = _resolve_close(disp_name, df)

        if yf_prev is not None and curr_close is not None:
            weekly_pct = calculate_weekly_pct(yf_prev, curr_close)
            data_quality["indices_ok"] += 1
        else:
            weekly_pct = None
            src = "incomplete"

        close_rounded = round2(curr_close)
        resolved_closes[disp_name] = curr_close  # store for EMA/S&R use
        print(f"  {disp_name:15s}: close={close_rounded}  prev={round2(yf_prev)}  weekly%={weekly_pct}  [{src}]")
        indices.append({"name": disp_name, "close": close_rounded, "pct": weekly_pct})

    # Build sectors data
    print("\n" + "-" * 70)
    print("PROCESSING SECTORS")
    print("-" * 70)

    sectors = []
    data_quality["sectors_ok"] = 0
    data_quality["sectors_total"] = len(SECTOR_NAMES)

    for disp_name in SECTOR_NAMES:
        df = yf_cache.get(disp_name)
        yf_prev, curr_close, src = _resolve_close(disp_name, df)

        if yf_prev is not None and curr_close is not None:
            weekly_pct = calculate_weekly_pct(yf_prev, curr_close)
            data_quality["sectors_ok"] += 1
        else:
            weekly_pct = None
            src = "incomplete"

        close_rounded = round2(curr_close)
        print(f"  {disp_name:20s}: close={close_rounded}  prev={round2(yf_prev)}  weekly%={weekly_pct}  [{src}]")
        sectors.append({"name": disp_name, "close": close_rounded, "pct": weekly_pct})

    # Build EMA data
    print("\n" + "-" * 70)
    print("COMPUTING EMAs")
    print("-" * 70)

    ema_to_display = {
        "NIFTY": "NIFTY 50",
        "BANK NIFTY": "BANK NIFTY",
        "FINNIFTY": "FINNIFTY"
    }

    ema_rows = []

    for ema_name in EMA_NAMES:
        disp_name = ema_to_display[ema_name]
        df = yf_cache.get(disp_name)

        if df is None or df.empty or "Close" not in df.columns:
            print(f"  [WARN] {ema_name}: No data for EMA calculation")
            ema_rows.append({
                "name": ema_name,
                "close": None,
                "emas": {},
                "bias": "VOLATILE"
            })
            continue

        try:
            ts_curr = pd.Timestamp(curr_friday_actual)
            valid_df = df.dropna(subset=["Close"])
            close_series = valid_df["Close"]

            # Use the resolved close (NSE-corrected) for display and bias,
            # not the raw YF last close which may be stale (e.g. FINNIFTY Jul 17).
            resolved_curr = resolved_closes.get(disp_name)
            if resolved_curr is not None:
                close_val = round0(float(resolved_curr))
                
                # ── NEW: Patch the missing Close so EMAs are calculated on the correct final day ──
                if ts_curr not in close_series.index or pd.isna(close_series.get(ts_curr)):
                    close_series.loc[ts_curr] = resolved_curr
                    close_series = close_series.sort_index()
            else:
                subset = valid_df[valid_df.index <= ts_curr]
                if subset.empty:
                    raise ValueError(f"No data on or before {curr_friday_actual}")
                close_val = round0(float(subset["Close"].iloc[-1]))

            # Calculate all EMAs (uses full YF history for accuracy)
            emas = {}
            for period in EMA_PERIODS:
                ema_series = calculate_ema(close_series, period)
                ema_subset = ema_series[ema_series.index <= ts_curr]

                if not ema_subset.empty:
                    emas[period] = round0(float(ema_subset.iloc[-1]))
                else:
                    emas[period] = None

            # Determine bias from EMA-200
            bias = determine_bias(close_val, emas.get(200))

            ema_rows.append({
                "name": ema_name,
                "close": close_val,
                "emas": emas,
                "bias": bias
            })

            print(f"  {ema_name:12s}: close={close_val:,}  bias={bias}  "
                  f"EMAs={list(emas.values())}")

        except Exception as e:
            print(f"  [ERROR] {ema_name}: EMA calculation failed: {e}")
            ema_rows.append({
                "name": ema_name,
                "close": None,
                "emas": {},
                "bias": "VOLATILE"
            })

    # Build Support/Resistance data
    print("\n" + "-" * 70)
    print("COMPUTING SUPPORT/RESISTANCE")
    print("-" * 70)

    sr_map = {
        "NIFTY 50": "NIFTY",
        "BANK NIFTY": "BANK NIFTY",
        "FINNIFTY": "FINNIFTY"
    }

    sr_rows = []

    for disp_name, ema_name in sr_map.items():
        ema_row = next((e for e in ema_rows if e["name"] == ema_name), None)
        bias = ema_row["bias"] if ema_row else "VOLATILE"
        close = ema_row["close"] if ema_row else None

        # Get rounding rule (index-specific, not threshold-based)
        rounding_multiple = SR_ROUNDING_RULES.get(ema_name, 100)

        if close is None:
            print(f"  [WARN] {ema_name}: No close for S/R calculation")
            sr_rows.append({
                "name": ema_name,
                "bias": bias,
                "close": close,
                "s1": None,
                "s2": None,
                "r1": None,
                "r2": None
            })
            continue

        # Per-index tuned factors
        factors = SR_FACTORS.get(ema_name, {})
        sr_levels = calculate_support_resistance(
            close, rounding_multiple,
            s1_factor=factors.get("s1", 0.98),
            s2_factor=factors.get("s2", 0.967),
            r1_factor=factors.get("r1", 1.02),
            r2_factor=factors.get("r2", 1.033),
        )

        sr_rows.append({
            "name": ema_name,
            "bias": bias,
            "close": close,
            **sr_levels
        })

        print(f"  {ema_name:12s}: S2={sr_levels['s2']:,}  S1={sr_levels['s1']:,}  "
              f"C={close:,}  R1={sr_levels['r1']:,}  R2={sr_levels['r2']:,}")

    return {
        "indices": indices,
        "sectors": sectors,
        "ema": ema_rows,
        "sr": sr_rows,
        "data_quality": data_quality,
        "actual_fridays": (prev_friday_actual, curr_friday_actual),
        "target_fridays": (prev_friday_target, curr_friday_target),
        "yf_cache": yf_cache
    }


def build_narrative(mkt_data, fii_dii, constituents, global_mkts, start_date, end_date, derivatives_data=None):
    """
    Generate all narrative text for the report.

    Returns dict with keys: wgb (list of 6 paragraphs), deriv (list of 2 paragraphs)
    """
    idx_dict = {d["name"]: d for d in mkt_data["indices"]}
    sectors = mkt_data["sectors"]

    def _close(name):
        v = idx_dict.get(name)
        return v.get("close") if v else None

    def _pct(name):
        v = idx_dict.get(name)
        return v.get("pct") if v else None

    def _fmt_close(name):
        c = _close(name)
        return "N/A" if c is None else f"{round0(c):,}"

    def _fmt_pct(name):
        p = _pct(name)
        if p is None:
            return "data unavailable"
        return f"{abs(p):.2f}% {'gain' if p >= 0 else 'cut'}"

    # Paragraph 1: Benchmark + Banking
    nifty_pct_val = _pct("NIFTY 50") or 0
    bank_pct_val = _pct("BANK NIFTY") or 0
    nifty_bias = "positive" if nifty_pct_val >= 0 else "negative"
    
    # Determine Bank Nifty relative performance
    diff = bank_pct_val - nifty_pct_val
    if diff > 0.1:
        bank_perf = "outperformed"
    elif diff < -0.1:
        bank_perf = "underperformed"
    else:
        bank_perf = "performed in line with"

    p1 = (f"Benchmark Index traded with {nifty_bias} bias in the previous week "
          f"before closing with {_fmt_pct('NIFTY 50')} at {_fmt_close('NIFTY 50')} level. "
          f"Banking Index {bank_perf} the benchmark index as it "
          f"closed with {_fmt_pct('BANK NIFTY')} at {_fmt_close('BANK NIFTY')} level.")

    # Paragraph 2: Constituents
    if constituents:
        ng = constituents.get("nifty_gainers", [])
        nl = constituents.get("nifty_losers", [])
        bg = constituents.get("banknifty_gainers", [])[:1]
        bl = constituents.get("banknifty_losers", [])[:1]

        def _names(items):
            return " and ".join(d["name"] for d in items)

        def _pcts(items):
            return " and ".join(f"{abs(d['pct']):.2f}%" for d in items)

        if ng and nl:
            n_txt = (f"Among the Nifty constituents, {_names(ng)} outperformed "
                     f"the benchmark index as they closed with {_pcts(ng)} gains "
                     f"while {_names(nl)} underperformed as they closed with "
                     f"{_pcts(nl)} cut respectively.")
        else:
            n_txt = ("Among the Nifty constituents, please refer to NSE for "
                     "top gainers and losers for the week.")

        if bg and bl:
            b_txt = (f" Among the Bank Nifty constituents {_names(bg)} "
                     f"outperformed the benchmark index as it closed with {_pcts(bg)} gain "
                     f"while {_names(bl)} underperformed the benchmark index as it closed with "
                     f"{_pcts(bl)} cut.")
        else:
            b_txt = (" Among the Bank Nifty constituents, please refer to NSE "
                     "for weekly top performers and laggards.")

        p2 = n_txt + b_txt
    else:
        p2 = ("Among the Nifty constituents, please refer to NSE for top gainers "
              "and losers for the week. Among the Bank Nifty constituents, please "
              "refer to NSE for weekly top performers and laggards.")

    # Paragraph 3: Sectors + Broader market
    valid_sectors = [s for s in sectors if s.get("pct") is not None]
    broader_names = {"NIFTYMIDCA", "NIFTYSMLC", "MIDCAP SELECT"}
    pure_sectors = [s for s in valid_sectors if s["name"] not in broader_names]

    if pure_sectors:
        top_gainer = max(pure_sectors, key=lambda x: x["pct"])
        top_loser = min(pure_sectors, key=lambda x: x["pct"])

        gainer_word = "gain" if top_gainer['pct'] >= 0 else "cut"
        loser_word = "gain" if top_loser['pct'] >= 0 else "cut"

        p3 = (f"Among sectors, {top_gainer['name'].replace(' ', '')} index outperformed the benchmark "
              f"index during the previous week as it closed with {abs(top_gainer['pct']):.2f}% {gainer_word} while "
              f"{top_loser['name'].replace(' ', '')} index underperformed as it closed with "
              f"{abs(top_loser['pct']):.2f}% {loser_word}.")
    else:
        p3 = "Sector performance data is available on NSE."

    # Add broader market commentary
    mid_sec = next((s for s in valid_sectors if "MIDCA" in s["name"].upper() or
                    "MIDCAP SELECT" in s["name"]), None)
    small_sec = next((s for s in valid_sectors if "SMLC" in s["name"].upper()), None)

    if mid_sec and small_sec:
        mid_pct = mid_sec["pct"]
        small_pct = small_sec["pct"]
        nifty_pct = _pct("NIFTY 50") or 0
        
        mid_dir = "gain" if mid_pct >= 0 else "cut"
        small_dir = "gain" if small_pct >= 0 else "cut"
        
        # Helper to determine relative performance word
        def get_rel_perf(pct, bm_pct):
            diff = pct - bm_pct
            if diff > 0.1:
                return "outperformed"
            elif diff < -0.1:
                return "lagged" if pct >= 0 else "underperformed"
            else:
                return "performed in line with"

        mid_rel = get_rel_perf(mid_pct, nifty_pct)
        small_rel = get_rel_perf(small_pct, nifty_pct)
        
        if mid_rel == small_rel and mid_dir == small_dir:
            p3 += (f" Broader market {mid_rel} the benchmark index as both MIDCAP index "
                   f"and SMALLCAP index closed with {abs(mid_pct):.2f}% and {abs(small_pct):.2f}% "
                   f"{mid_dir}s respectively.")
        else:
            p3 += (f" MIDCAP index {mid_rel} the benchmark index as it closed with {abs(mid_pct):.2f}% {mid_dir} "
                   f"while SMALLCAP index {small_rel} the benchmark index as it closed with {abs(small_pct):.2f}% {small_dir}.")

    # Paragraph 4: VIX
    vix_pct = _pct("VIX")
    vix_close = _close("VIX")

    if vix_pct is not None and vix_close is not None:
        p4 = (f"Volatility index (India VIX) closed with {abs(vix_pct):.2f}% "
              f"{'cut' if vix_pct < 0 else 'gain'} at {vix_close:.2f} level.")
    else:
        p4 = "Volatility index (India VIX) data unavailable."

    # Paragraph 5: FII/DII
    fii_val = fii_dii.get("fii")
    dii_val = fii_dii.get("dii")

    if fii_val is not None and dii_val is not None:
        fii_action = "buyers" if fii_val >= 0 else "sellers"
        dii_action = "buyers" if dii_val >= 0 else "sellers"
        fii_verb = "bought" if fii_val >= 0 else "sold"
        dii_verb = "bought" if dii_val >= 0 else "sold"

        is_weekly = fii_dii.get("is_weekly", False)
        is_complete = fii_dii.get("is_complete", False)
        days_covered = fii_dii.get("days_covered", 0)

        if is_weekly and is_complete:
            time_period_str = "during the previous week in the cash segment."
        elif is_weekly and not is_complete:
            time_period_str = (
                f"over the {days_covered} trading day(s) logged so far this week "
                f"in the cash segment (full week total will complete automatically "
                f"as the remaining days are logged)."
            )
        else:
            time_period_str = "on the latest available trading day (weekly log not yet available for this week)."

        if fii_action == dii_action:
            p5 = (f"Both FII and DII were net {fii_action} as they {fii_verb} stocks worth "
                  f"Rs {abs(fii_val):,} Cr and {abs(dii_val):,} Cr respectively {time_period_str}")
        else:
            p5 = (f"FII were net {fii_action} as they {fii_verb} stocks worth Rs {abs(fii_val):,} Cr while "
                  f"DII were net {dii_action} as they {dii_verb} stocks worth Rs {abs(dii_val):,} Cr {time_period_str}")
    else:
        p5 = "FII/DII data unavailable. Please refer to NSE/SEBI for latest figures."

    # Paragraph 6: Global markets
    if global_mkts:
        djia_pct = global_mkts.get("djia")
        stoxx_pct = global_mkts.get("stoxx")

        if djia_pct is not None and stoxx_pct is not None:
            djia_dir = "gains" if djia_pct >= 0 else "cuts"
            stoxx_dir = "gains" if stoxx_pct >= 0 else "cuts"
            djia_bias = "positive" if djia_pct >= 0 else "negative"
            stoxx_bias = "positive" if stoxx_pct >= 0 else "negative"
            overall_bias = "a positive bias" if djia_bias == stoxx_bias and djia_bias == "positive" else \
                           "a negative bias" if djia_bias == stoxx_bias else "a mixed bias"

            if djia_dir == stoxx_dir:
                p6 = (f"Globally, equity markets traded with {overall_bias} as both US market (DJIA) "
                      f"and European market (STOXX 600) closed with {abs(djia_pct):.2f}% and {abs(stoxx_pct):.2f}% {djia_dir} respectively.")
            else:
                p6 = (f"Globally, equity markets traded with {overall_bias} as US market (DJIA) "
                      f"closed with {abs(djia_pct):.2f}% {djia_dir} while European market (STOXX 600) "
                      f"closed with {abs(stoxx_pct):.2f}% {stoxx_dir}.")
        elif djia_pct is not None:
            p6 = (f"Globally, US market (DJIA) closed with {abs(djia_pct):.2f}% "
                  f"{'gain' if djia_pct >= 0 else 'cut'}. Please refer to Bloomberg "
                  f"for European market (STOXX 600) movements.")
        else:
            p6 = ("Globally, equity markets performance for the week: please refer to "
                  "Bloomberg or Reuters for US (DJIA) and European (STOXX 600) market movements.")
    else:
        p6 = ("Globally, equity markets performance for the week: please refer to "
              "Bloomberg or Reuters for US (DJIA) and European (STOXX 600) market movements.")

    # Derivatives paragraphs
    # Use the actual friday date for week ending
    _, curr_friday_actual = mkt_data["actual_fridays"]
    week_end_str = curr_friday_actual.strftime("%B %d, %Y")
    
    derivatives_data = derivatives_data or {}
    
    def _build_deriv_para(idx_key, display_name):
        data = derivatives_data.get(idx_key)
        if data and all(data.values()):
            pcr = float(data["pcr"])
            
            if data['call_oi'] == data['put_oi']:
                max_oi_str = f"max OI at {data['call_oi']} for both call and put side"
            else:
                max_oi_str = f"max OI at {data['call_oi']} call and {data['put_oi']} put side"
                
            if data['oi_add_c'] == data['oi_add_p']:
                max_oi_add_str = f"max OI addition at {data['oi_add_c']} for both call and put side"
            else:
                max_oi_add_str = f"max OI addition at {data['oi_add_c']} call and {data['oi_add_p']} put side"
                
            return (f"Derivatives data for {display_name} suggest {max_oi_add_str} with {max_oi_str} "
                    f"({data['expiry']} expiry). "
                    f"{display_name} cumulative PCR closed at {pcr:.2f} for {data['pcr_date']}.")
        else:
            return (f"Derivatives data for {display_name}: please refer to NSE for current OI distribution, "
                    f"max OI call/put strikes, and PCR for the upcoming expiry.")

    d1 = _build_deriv_para("NIFTY", "Nifty")
    d2 = _build_deriv_para("BANK NIFTY", "Bank-nifty")

    return {
        "wgb": [p1, p2, p3, p4, p5, p6],
        "deriv": [d1, d2]
    }


def _build_inline_image_xml(rId, cx, cy, name, desc):
    """
    Build the XML for an inline image with exact EMU dimensions.
    Uses raw XML string + lxml.fromstring() to handle DrawingML namespaces correctly
    (lxml rejects xmlns:* as regular attributes, so we parse from string instead).
    """
    from lxml import etree
    xml_str = (
        '<w:drawing'
        ' xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        ' xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"'
        ' xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"'
        ' xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture"'
        ' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"'
        '>'
        '<wp:inline distT="0" distB="0" distL="0" distR="0">'
        f'<wp:extent cx="{cx}" cy="{cy}"/>'
        '<wp:effectExtent l="0" t="0" r="0" b="0"/>'
        f'<wp:docPr id="1" name="{name}" descr="{desc}"/>'
        '<wp:cNvGraphicFramePr>'
        '<a:graphicFrameLocks noChangeAspect="1"/>'
        '</wp:cNvGraphicFramePr>'
        '<a:graphic>'
        '<a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">'
        '<pic:pic>'
        '<pic:nvPicPr>'
        f'<pic:cNvPr id="1" name="{name}"/>'
        '<pic:cNvPicPr/>'
        '</pic:nvPicPr>'
        '<pic:blipFill>'
        f'<a:blip r:embed="{rId}"/>'
        '<a:stretch><a:fillRect/></a:stretch>'
        '</pic:blipFill>'
        '<pic:spPr>'
        '<a:xfrm><a:off x="0" y="0"/>'
        f'<a:ext cx="{cx}" cy="{cy}"/>'
        '</a:xfrm>'
        '<a:prstGeom prst="rect"><a:avLst/></a:prstGeom>'
        '</pic:spPr>'
        '</pic:pic>'
        '</a:graphicData>'
        '</a:graphic>'
        '</wp:inline>'
        '</w:drawing>'
    )
    return etree.fromstring(xml_str)


def fill_docx_document(mkt_data, fii_dii, narrative, start_date, end_date,
                       output_path, tech_outlook=None, chart_paths=None):
    """
    Fill the template DOCX with all data.

    This function handles:
    Page 1:
    - Header date patching
    - Week Gone By narrative
    - Derivatives narrative
    - Indices table
    - Sectors table
    - FII/DII table
    - Support/Resistance table
    - EMA table
    Page 2:
    - TECHNICAL OUTLOOK heading
    - Chart + Commentary table (3 rows: NIFTY, BANKNIFTY, FINNIFTY)
    - EQUITY RESEARCH TEAM heading
    - Contact table
    - Disclaimer
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from lxml import etree
    import copy
    import math
    import re
    import shutil
    import zipfile

    print("\n" + "=" * 70)
    print("GENERATING DOCX REPORT")
    print("=" * 70)

    # Step 1: Copy template and prepare header date string
    # NOTE: Header zip-patching is done AFTER doc.save() to avoid python-docx
    # overwriting the gradient header design when it saves the document.
    print("\n[1/6] Preparing document...")
    _, curr_friday_actual = mkt_data["actual_fridays"]
    # Header uses Monday (End Sunday + 1 day) per user request
    header_monday = end_date + timedelta(days=1)
    # Use non-zero-padded day to match reference format: "July 27, 2026" not "July 027, 2026"
    header_date_str = header_monday.strftime("%B %-d, %Y") if os.name != 'nt' else \
                      header_monday.strftime("%B %#d, %Y")

    shutil.copy2(TEMPLATE_PATH, output_path)

    # Step 2: Load document
    print("\n[2/6] Loading document structure...")
    doc = Document(output_path)

    # Helper: replace paragraph text
    def _replace_para_text(para, new_text, bold=None, color_hex=None, font_size=None):
        """Replace all runs in paragraph with single run."""
        p_elem = para._p
        existing = p_elem.findall(qn('w:r'))
        base_rPr = None

        if existing:
            first_rPr = existing[0].find(qn('w:rPr'))
            if first_rPr is not None:
                base_rPr = copy.deepcopy(first_rPr)

        # Remove existing runs
        for child in list(p_elem):
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            if tag in ('r', 'hyperlink', 'ins', 'del'):
                p_elem.remove(child)

        # Create new run properties
        rPr = copy.deepcopy(base_rPr) if base_rPr is not None else etree.Element(qn('w:rPr'))

        if bold is not None:
            b = rPr.find(qn('w:b'))
            bCs = rPr.find(qn('w:bCs'))
            if bold:
                if b is None:
                    etree.SubElement(rPr, qn('w:b'))
                if bCs is None:
                    etree.SubElement(rPr, qn('w:bCs'))
            else:
                for el in [b, bCs]:
                    if el is not None:
                        rPr.remove(el)

        if color_hex is not None:
            c = rPr.find(qn('w:color'))
            if c is None:
                c = etree.SubElement(rPr, qn('w:color'))
            c.set(qn('w:val'), color_hex)
            if c.get(qn('w:themeColor')):
                del c.attrib[qn('w:themeColor')]

        if font_size is not None:
            # Word uses half-points
            half_pts = str(int(font_size * 2))
            sz = rPr.find(qn('w:sz'))
            if sz is None:
                sz = etree.SubElement(rPr, qn('w:sz'))
            sz.set(qn('w:val'), half_pts)
            szCs = rPr.find(qn('w:szCs'))
            if szCs is None:
                szCs = etree.SubElement(rPr, qn('w:szCs'))
            szCs.set(qn('w:val'), half_pts)

        # Create new run
        r = etree.SubElement(p_elem, qn('w:r'))
        r.insert(0, rPr)
        t = etree.SubElement(r, qn('w:t'))
        t.text = new_text
        if new_text != new_text.strip():
            t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

    # Step 3: Fill narrative
    print("\n[3/6] Filling narrative content...")

    main_table = doc.tables[0]
    left_cell = main_table.rows[0].cells[0]
    right_cell = main_table.rows[0].cells[2]

    left_paras = left_cell.paragraphs
    wgb = narrative["wgb"]
    deriv = narrative["deriv"]

    def _tighten_para_spacing(para, before=40, after=60, line=216):
        """Override MBullet style spacing so all 8 narrative paras fit on Page 1."""
        p_elem = para._p
        pPr = p_elem.find(qn('w:pPr'))
        if pPr is None:
            pPr = etree.SubElement(p_elem, qn('w:pPr'))
            p_elem.insert(0, pPr)
        spacing = pPr.find(qn('w:spacing'))
        if spacing is None:
            spacing = etree.SubElement(pPr, qn('w:spacing'))
        spacing.set(qn('w:before'), str(before))
        spacing.set(qn('w:after'),  str(after))
        spacing.set(qn('w:line'),   str(line))
        spacing.set(qn('w:lineRule'), 'auto')
        ctxSpc = pPr.find(qn('w:contextualSpacing'))
        if ctxSpc is None:
            ctxSpc = etree.SubElement(pPr, qn('w:contextualSpacing'))
        ctxSpc.set(qn('w:val'), '1')

    # Find "WEEK GONE BY" heading and tighten it
    wgb_start = 0
    for i, p in enumerate(left_paras):
        if "week gone by" in p.text.strip().lower():
            wgb_start = i
            _tighten_para_spacing(left_paras[i], before=0, after=20)
            break

    # Write WGB bullets — spread vertically to fill the full column height
    # before=80  ≈ 4pt above each bullet
    # after=220  ≈ 11pt below each bullet  → distributes empty column space evenly
    for i, text in enumerate(wgb):
        idx = wgb_start + 1 + i
        if idx < len(left_paras):
            _replace_para_text(left_paras[idx], text, font_size=9.0)
            # Removed _tighten_para_spacing to allow template's native spacing
        else:
            print(f"  [WARN] WGB para {i} out of bounds")

    # Find "Derivatives" heading — give extra space before so it visually separates from WGB
    deriv_start = wgb_start + len(wgb) + 1
    for i in range(deriv_start, len(left_paras)):
        if left_paras[i].text.strip().lower() == "derivatives":
            deriv_start = i
            # Removed _tighten_para_spacing to allow template's native spacing
            break

    # Write derivatives bullets — same vertical spread as WGB
    for i, text in enumerate(deriv):
        idx = deriv_start + 1 + i
        if idx < len(left_paras):
            _replace_para_text(left_paras[idx], text, font_size=9.0)
            # Removed _tighten_para_spacing to allow template's native spacing
        else:
            print(f"  [WARN] Deriv para {i} out of bounds")

    # Step 4: Fill tables
    print("\n[4/6] Filling Page 1 tables...")

    right_tables = right_cell.tables

    # Helper: find table by keywords
    def _find_table(tables, keywords):
        kws = [k.lower() for k in keywords]
        for t in tables:
            header = " ".join(c.text for c in t.rows[0].cells).lower()
            if all(k in header for k in kws):
                return t
        return None

    # Helper: set cell text
    def _set_cell(cell, text, para_idx=0, bold=None, color_hex=None, font_size=None, align=None):
        if para_idx < len(cell.paragraphs):
            _replace_para_text(cell.paragraphs[para_idx], text, bold=bold, color_hex=color_hex, font_size=font_size)
            if align is not None:
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                cell.paragraphs[para_idx].alignment = align

    def _apply_clean_table_borders(table):
        """Apply green top/bottom borders, remove vertical/inside borders, bold/center header."""
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = etree.SubElement(tbl, qn('w:tblPr'))
        
        tblBorders = tblPr.find(qn('w:tblBorders'))
        if tblBorders is not None:
            tblPr.remove(tblBorders)
        
        tblBorders = etree.SubElement(tblPr, qn('w:tblBorders'))
        
        # Table top border (green)
        top = etree.SubElement(tblBorders, qn('w:top'))
        top.set(qn('w:val'), 'single')
        top.set(qn('w:sz'), '12')
        top.set(qn('w:color'), COLOR_GREEN)
        
        # Table bottom border (green)
        bottom = etree.SubElement(tblBorders, qn('w:bottom'))
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '12')
        bottom.set(qn('w:color'), COLOR_GREEN)
        
        # Remove inner and vertical borders
        for border_name in ['left', 'right', 'insideV', 'insideH']:
            b = etree.SubElement(tblBorders, qn(f'w:{border_name}'))
            b.set(qn('w:val'), 'none')
            
        # Add bottom border to header row cells
        if len(table.rows) > 0:
            for cell in table.rows[0].cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcBorders = tcPr.find(qn('w:tcBorders'))
                if tcBorders is None:
                    tcBorders = etree.SubElement(tcPr, qn('w:tcBorders'))
                cb_bottom = tcBorders.find(qn('w:bottom'))
                if cb_bottom is None:
                    cb_bottom = etree.SubElement(tcBorders, qn('w:bottom'))
                cb_bottom.set(qn('w:val'), 'double')
                cb_bottom.set(qn('w:sz'), '4')
                cb_bottom.set(qn('w:color'), '3C9114')

        # Style header row (bold, centered data columns)
        for i, cell in enumerate(table.rows[0].cells):
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.bold = True
                if i > 0:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Guard: check if val is a real number (not None, not NaN)
    def _is_valid(val):
        """Return True if val is a real number (not None, not NaN)."""
        if val is None:
            return False
        try:
            return not math.isnan(float(val))
        except (TypeError, ValueError):
            return False

    # Helper: percentage color
    def _pct_color(val):
        if not _is_valid(val):
            return None
        if val > 0:
            return COLOR_GREEN
        elif val < 0:
            return COLOR_RED
        return "000000"  # Black for exactly 0

    # Indices table
    idx_table = _find_table(right_tables, ["index", "close"]) or right_tables[0]
    for i, row_data in enumerate(mkt_data["indices"], 1):
        if i >= len(idx_table.rows):
            break
        row = idx_table.rows[i]
        name = row_data["name"]
        close = row_data["close"]
        pct = row_data["pct"]

        # Format close (VIX uses 2 decimals, others use comma format)
        if name == "VIX":
            close_str = f"{close:.2f}" if _is_valid(close) else ""
        else:
            close_str = f"{round0(close):,}" if _is_valid(close) else ""

        pct_str = f"{pct:.2f}" if _is_valid(pct) else ""

        formatted_name = name.replace(" ", "")
        _set_cell(row.cells[0], formatted_name, bold=True)
        _set_cell(row.cells[1], close_str, align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell(row.cells[2], pct_str, color_hex=_pct_color(pct), align=WD_ALIGN_PARAGRAPH.CENTER)

    _apply_clean_table_borders(idx_table)
    print(f"  [OK] Indices table: {len(mkt_data['indices'])} rows filled")

    # Sectors table
    sec_table = _find_table(right_tables, ["sector"]) or right_tables[1]
    for i, row_data in enumerate(mkt_data["sectors"], 1):
        if i >= len(sec_table.rows):
            break
        row = sec_table.rows[i]
        name = row_data["name"]
        close = row_data["close"]
        pct = row_data["pct"]

        close_str = f"{round0(close):,}" if _is_valid(close) else ""
        pct_str = f"{pct:.2f}" if _is_valid(pct) else ""

        formatted_name = name.replace(" ", "")
        _set_cell(row.cells[0], formatted_name, bold=True)
        _set_cell(row.cells[1], close_str, align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell(row.cells[2], pct_str, color_hex=_pct_color(pct), align=WD_ALIGN_PARAGRAPH.CENTER)

    _apply_clean_table_borders(sec_table)
    print(f"  [OK] Sectors table: {len(mkt_data['sectors'])} rows filled")

    # FII/DII table
    fii_table = _find_table(right_tables, ["rs in cr"]) or right_tables[2]
    fii_val = fii_dii.get("fii")
    dii_val = fii_dii.get("dii")

    if len(fii_table.rows) >= 3:
        fii_str = f"{fii_val:,}" if fii_val is not None else ""
        dii_str = f"{dii_val:,}" if dii_val is not None else ""

        _set_cell(fii_table.rows[1].cells[1], fii_str,
                  color_hex=COLOR_GREEN if fii_val is not None and fii_val >= 0 else COLOR_RED)
        _set_cell(fii_table.rows[2].cells[1], dii_str,
                  color_hex=COLOR_GREEN if dii_val is not None and dii_val >= 0 else COLOR_RED)

    _apply_clean_table_borders(fii_table)
    print(f"  [OK] FII/DII table filled")

    # Support/Resistance table
    sr_table = _find_table(doc.tables, ["support", "resistance"]) or doc.tables[1]
    sr_data_start = 2

    for i, sr_row in enumerate(mkt_data["sr"]):
        row_idx = sr_data_start + i
        if row_idx >= len(sr_table.rows):
            break

        row = sr_table.rows[row_idx]
        _set_cell(row.cells[0], sr_row["name"].replace(" ", ""), bold=True)
        _set_cell(row.cells[1], sr_row["bias"], bold=True)
        _set_cell(row.cells[2], f"{sr_row['s2']:,}" if sr_row['s2'] is not None else "")
        _set_cell(row.cells[3], f"{sr_row['s1']:,}" if sr_row['s1'] is not None else "")
        _set_cell(row.cells[4], f"{sr_row['close']:,}" if sr_row['close'] is not None else "")
        _set_cell(row.cells[5], f"{sr_row['r1']:,}" if sr_row['r1'] is not None else "")
        _set_cell(row.cells[6], f"{sr_row['r2']:,}" if sr_row['r2'] is not None else "")

    # Clear any extra template rows beyond what we filled
    filled_up_to = sr_data_start + len(mkt_data["sr"])
    for extra_idx in range(filled_up_to, len(sr_table.rows)):
        extra_row = sr_table.rows[extra_idx]
        for cell in extra_row.cells:
            _set_cell(cell, "")

    print(f"  [OK] S/R table: {len(mkt_data['sr'])} rows filled")


    # EMA table
    ema_table = _find_table(doc.tables, ["ema"]) or doc.tables[2]
    
    # Format the EMA header to match the S/R header style
    _set_cell(ema_table.rows[0].cells[0], "DAILY EXPONENTIAL MOVING AVERAGE", 
              bold=True, color_hex=HEADING_COLOR, font_size=11)
              
    ema_data_start = 2

    for i, ema_row in enumerate(mkt_data["ema"]):
        row_idx = ema_data_start + i
        if row_idx >= len(ema_table.rows):
            break

        row = ema_table.rows[row_idx]
        _set_cell(row.cells[0], ema_row["name"].replace(" ", ""), bold=True)

        close_str = f"{ema_row['close']:,}" if ema_row['close'] is not None else ""
        _set_cell(row.cells[1], close_str)

        for col_idx, period in enumerate(EMA_PERIODS, 2):
            ema_val = ema_row["emas"].get(period)
            ema_str = f"{ema_val:,}" if ema_val is not None else ""
            _set_cell(row.cells[col_idx], ema_str)

    print(f"  [OK] EMA table: {len(mkt_data['ema'])} rows filled")

    # ═══════════════════════════════════════════════════════════════
    # STEP 5/6: Build Page 2 — Technical Outlook
    # ═══════════════════════════════════════════════════════════════
    print("\n[5/6] Building Page 2 — Technical Outlook...")

    body = doc.element.body

    # The final sectPr is the very last child of body (always)
    final_sectPr = None
    for child in reversed(list(body)):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'sectPr':
            final_sectPr = child
            break

    if final_sectPr is None:
        print("  [WARN] No final sectPr found — document structure may be abnormal")

    # ─── Remove ALL trailing paragraphs and inline sectPr paragraphs ───
    # The template contains:
    #   [5]  empty <p> between S/R and EMA tables
    #   [7]  <p> with inline <w:sectPr> — this is a spurious section break that
    #         pushes the EMA table onto page 2.
    #   [8]  <bookmarkStart>
    #   [9]  <bookmarkEnd>
    #   [10] trailing empty <p>
    # We remove all <p> and bookmark elements that come AFTER the last table,
    # so that only our new pb_para creates the page break after the EMA table.

    def _get_tag(el):
        return el.tag.split('}')[-1] if '}' in el.tag else el.tag

    def _para_has_sectpr(p_elem):
        """Return True if this <w:p> contains a nested <w:sectPr>."""
        for child in p_elem:
            t = _get_tag(child)
            if t == 'pPr':
                for grandchild in child:
                    if _get_tag(grandchild) == 'sectPr':
                        return True
        return False

    import copy
    captured_page1_sectPr = None

    to_remove = []
    for child in reversed(list(body)):
        if child == final_sectPr:
            continue
        tag = _get_tag(child)
        if tag == 'p':
            if _para_has_sectpr(child) and captured_page1_sectPr is None:
                pPr_el = child.find(qn('w:pPr'))
                for grandchild in pPr_el:
                    if _get_tag(grandchild) == 'sectPr':
                        captured_page1_sectPr = copy.deepcopy(grandchild)
                        break
            to_remove.append(child)   # Remove all trailing <p> including inline sectPr ones
        elif tag in ('bookmarkStart', 'bookmarkEnd'):
            to_remove.append(child)   # Remove bookmarks too
        else:
            # We've hit a table (EMA table) — stop here
            break

    for child in to_remove:
        body.remove(child)

    # ─── Remove artifact paragraphs BETWEEN the main table and S/R table ───
    # Template body has a '2\t' artifact paragraph [3] and possibly other 
    # empty paragraphs between the main WEEK GONE BY table and the S/R table.
    # These create blank space on page 1. Remove any <p> nodes between tables.
    body_children = list(body)
    # Only remove the inter-table artifact paragraphs (not trailing ones already removed)
    # We'll specifically target the '2' artifact and blank spacers between tables
    for child in list(body):
        tag = _get_tag(child)
        if tag == 'p' and child not in body_children:
            continue  # skip newly inserted elements
        if tag == 'p' and not _para_has_sectpr(child):
            # Check if text is just digits/whitespace (artifact like '2\t')
            texts = ''.join(
                t.text or '' for t in child.iter()
                if t.tag.endswith('}t')
            ).strip()
            if texts in ('', '2', '1', '3'):
                # Check it's between two tables
                idx = list(body).index(child)
                before_tags = [_get_tag(list(body)[j]) for j in range(max(0, idx-3), idx)]
                after_tags = [_get_tag(list(body)[j]) for j in range(idx+1, min(len(list(body)), idx+4))]
                if 'tbl' in before_tags and 'tbl' in after_tags:
                    body.remove(child)

    # ─── POST-CLEANUP: Insert spacer between S/R and EMA tables ───
    # Must be done AFTER cleanup so it doesn't get removed by artifact logic.
    spacer_para = etree.Element(qn('w:p'))
    spacer_pPr = etree.SubElement(spacer_para, qn('w:pPr'))
    spacer_spc = etree.SubElement(spacer_pPr, qn('w:spacing'))
    spacer_spc.set(qn('w:before'), '120')
    spacer_spc.set(qn('w:after'), '120')
    spacer_rPr = etree.SubElement(spacer_pPr, qn('w:rPr'))
    spacer_sz = etree.SubElement(spacer_rPr, qn('w:sz'))
    spacer_sz.set(qn('w:val'), '8')  # 4pt font — tiny spacer line
    ema_table._element.addprevious(spacer_para)
    print("  [OK] Spacer paragraph inserted between S/R and EMA tables (post-cleanup)")

    # ─── POST-CLEANUP: Ensure top margin before first table ───
    # Template has 2 empty <p> before the main table that create the gap between
    # the gradient header graphic and "WEEK GONE BY". Leave them as-is.
    # Only add a spacer if they were somehow removed (first child is tbl).
    first_body_child = list(body)[0]
    if _get_tag(first_body_child) == 'tbl':
        top_margin = etree.Element(qn('w:p'))
        top_pPr = etree.SubElement(top_margin, qn('w:pPr'))
        top_spc = etree.SubElement(top_pPr, qn('w:spacing'))
        top_spc.set(qn('w:before'), '0')
        top_spc.set(qn('w:after'), '200')
        top_spc.set(qn('w:line'), '240')
        top_spc.set(qn('w:lineRule'), 'auto')
        first_body_child.addprevious(top_margin)
        print("  [OK] Top margin paragraph inserted before first table")
    else:
        print("  [OK] Template top spacer paragraphs preserved")

    # ─── Insert a page-break paragraph that carries the Page 1 sectPr ───
    # This creates the section break between Page 1 and Page 2.
    # We move the final sectPr's content into a mid-body paragraph's sectPr,
    # then replace the final sectPr with Page 2 section settings.
    #
    # Strategy:
    # 1. Clone final_sectPr into a new <w:p><w:pPr><w:sectPr/></w:pPr></w:p>
    # 2. Insert that paragraph before the final sectPr
    # 3. Replace final_sectPr content with Page 2 margin settings
    #
    # This ensures Word sees two sections: Page 1 (original margins) + Page 2.

    import copy

    # Prefer the template's ORIGINAL Page 1 section properties captured
    # earlier (correct top/right/bottom/left/header/footer margins sized
    # for the gradient header banner). Only fall back to cloning
    # final_sectPr if the template didn't have a separate Page 1 section
    # (shouldn't normally happen, but keeps this from hard-failing).
    if captured_page1_sectPr is not None:
        page1_sectPr_clone = captured_page1_sectPr
        print("  [OK] Restored template's original Page 1 margins (top=2160 dxa gradient-header clearance)")
    else:
        page1_sectPr_clone = copy.deepcopy(final_sectPr)
        print("  [WARN] No captured Page 1 sectPr — falling back to Page 2 margins (may overlap header)")

    # Create the page-break paragraph
    pb_para = etree.Element(qn('w:p'))
    pb_pPr = etree.SubElement(pb_para, qn('w:pPr'))
    pb_pPr.append(page1_sectPr_clone)

    # ── CRITICAL: ensure the inline sectPr declares a next-page section break ──
    # Without <w:type w:val="nextPage"/>, Word treats this as a continuous
    # section break which does NOT create a new page, making Page 2 invisible.
    inline_sectPr = pb_pPr.find(qn('w:sectPr'))
    if inline_sectPr is not None:
        existing_type = inline_sectPr.find(qn('w:type'))
        if existing_type is None:
            # Insert <w:type w:val="nextPage"/> as first child of sectPr
            sect_type_el = etree.Element(qn('w:type'))
            sect_type_el.set(qn('w:val'), 'nextPage')
            inline_sectPr.insert(0, sect_type_el)
            print("  [OK] Added nextPage section type to pb_para sectPr")
        else:
            existing_type.set(qn('w:val'), 'nextPage')
            print("  [OK] Updated pb_para sectPr type to nextPage")

    # Insert the page-break paragraph just before the final sectPr
    final_sectPr_idx = list(body).index(final_sectPr)
    body.insert(final_sectPr_idx, pb_para)

    # Now update final_sectPr to Page 2 settings
    # Remove all existing children from final_sectPr
    for child in list(final_sectPr):
        final_sectPr.remove(child)

    # Set Page 2 margins (from reference document measurements)
    from config import PAGE2_LEFT_MARGIN, PAGE2_RIGHT_MARGIN, PAGE2_TOP_MARGIN, PAGE2_BOTTOM_MARGIN

    # Section type: nextPage ensures pages 2+ are on a fresh page
    sect_type2 = etree.SubElement(final_sectPr, qn('w:type'))
    sect_type2.set(qn('w:val'), 'nextPage')

    pgMar = etree.SubElement(final_sectPr, qn('w:pgMar'))
    pgMar.set(qn('w:top'), str(PAGE2_TOP_MARGIN // 635))       # EMU to dxa: /635
    pgMar.set(qn('w:right'), str(PAGE2_RIGHT_MARGIN // 635))
    pgMar.set(qn('w:bottom'), str(PAGE2_BOTTOM_MARGIN // 635))
    pgMar.set(qn('w:left'), str(PAGE2_LEFT_MARGIN // 635))
    pgMar.set(qn('w:header'), '120')
    pgMar.set(qn('w:footer'), '216')
    pgMar.set(qn('w:gutter'), '0')

    # Add page size (A4 portrait)
    pgSz = etree.SubElement(final_sectPr, qn('w:pgSz'))
    pgSz.set(qn('w:w'), '11906')   # A4 width in twips
    pgSz.set(qn('w:h'), '16838')   # A4 height in twips

    print("  [OK] Section break inserted between Page 1 and Page 2")

    # ─── Insert elements before final_sectPr ───

    def _insert_before_sectpr(element):
        """Insert an element just before the final sectPr."""
        if final_sectPr is not None:
            body.insert(list(body).index(final_sectPr), element)
        else:
            body.append(element)

    # ─── A) "TECHNICAL OUTLOOK" heading ───
    tech_heading = etree.SubElement(body, qn('w:p'))
    # Move before sectPr
    body.remove(tech_heading)
    _insert_before_sectpr(tech_heading)

    # Paragraph properties — MHeading2 style, spacing.before=0
    tech_pPr = etree.SubElement(tech_heading, qn('w:pPr'))
    pStyle = etree.SubElement(tech_pPr, qn('w:pStyle'))
    pStyle.set(qn('w:val'), 'MHeading2')
    spacing = etree.SubElement(tech_pPr, qn('w:spacing'))
    spacing.set(qn('w:before'), '0')

    # Run with text
    tech_run = etree.SubElement(tech_heading, qn('w:r'))
    tech_rPr = etree.SubElement(tech_run, qn('w:rPr'))
    sz = etree.SubElement(tech_rPr, qn('w:sz'))
    sz.set(qn('w:val'), str(HEADING_FONT_SIZE_HALFPTS))
    szCs = etree.SubElement(tech_rPr, qn('w:szCs'))
    szCs.set(qn('w:val'), str(HEADING_FONT_SIZE_HALFPTS))
    color = etree.SubElement(tech_rPr, qn('w:color'))
    color.set(qn('w:val'), HEADING_COLOR)
    tech_t = etree.SubElement(tech_run, qn('w:t'))
    tech_t.text = "TECHNICAL OUTLOOK"

    print("  [OK] 'TECHNICAL OUTLOOK' heading inserted")

    # ─── B) Chart + Commentary table (3 rows × 2 cols) ───
    if tech_outlook is not None and chart_paths is not None:
        # Build table XML manually for precise control
        tbl = etree.Element(qn('w:tbl'))
        _insert_before_sectpr(tbl)

        # Table properties
        tblPr = etree.SubElement(tbl, qn('w:tblPr'))

        # Table style (inherit from document)
        tblStyle = etree.SubElement(tblPr, qn('w:tblStyle'))
        tblStyle.set(qn('w:val'), 'TableGrid')
        
        tblW = etree.SubElement(tblPr, qn('w:tblW'))
        tblW.set(qn('w:w'), str(CHART_TABLE_WIDTH))
        tblW.set(qn('w:type'), 'dxa')

        # Cell margins — tight for maximum chart space
        tblCellMar = etree.SubElement(tblPr, qn('w:tblCellMar'))
        for mar_name in ['top', 'bottom']:
            m = etree.SubElement(tblCellMar, qn(f'w:{mar_name}'))
            m.set(qn('w:w'), '28')
            m.set(qn('w:type'), 'dxa')
        for mar_name in ['left', 'right']:
            m = etree.SubElement(tblCellMar, qn(f'w:{mar_name}'))
            m.set(qn('w:w'), '57')
            m.set(qn('w:type'), 'dxa')

        # Fixed layout
        tblLayout = etree.SubElement(tblPr, qn('w:tblLayout'))
        tblLayout.set(qn('w:type'), 'fixed')

        # Borders — single green lines on all sides
        tblBorders = etree.SubElement(tblPr, qn('w:tblBorders'))
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            b = etree.SubElement(tblBorders, qn(f'w:{border_name}'))
            b.set(qn('w:val'), 'single')
            b.set(qn('w:sz'), str(CHART_TABLE_BORDER_SZ))
            b.set(qn('w:color'), CHART_TABLE_BORDER_COLOR)
            b.set(qn('w:space'), '0')

        # Grid columns
        tblGrid = etree.SubElement(tbl, qn('w:tblGrid'))
        gridCol0 = etree.SubElement(tblGrid, qn('w:gridCol'))
        gridCol0.set(qn('w:w'), str(CHART_TABLE_COL0_WIDTH))
        gridCol1 = etree.SubElement(tblGrid, qn('w:gridCol'))
        gridCol1.set(qn('w:w'), str(CHART_TABLE_COL1_WIDTH))

        # Row data: chart name in order — keys MUST match generate_all_charts() output
        chart_order = ["NIFTY", "BANK NIFTY", "FINNIFTY"]
        display_headers = {
            "NIFTY":     "NIFTY (WEEKLY)",
            "BANK NIFTY": "BANKNIFTY (WEEKLY)",
            "FINNIFTY":  "FINNIFTY (WEEKLY)",
        }

        for row_idx, chart_name in enumerate(chart_order):
            tr = etree.SubElement(tbl, qn('w:tr'))

            # Row properties — set height as "atLeast" so rows expand to fit commentary
            trPr = etree.SubElement(tr, qn('w:trPr'))
            trHeight = etree.SubElement(trPr, qn('w:trHeight'))
            trHeight.set(qn('w:val'), str(CHART_TABLE_ROW_HEIGHTS[row_idx]))
            trHeight.set(qn('w:hRule'), 'atLeast')  # rows grow if content is taller

            # ─── Cell 0: Chart image ───
            tc0 = etree.SubElement(tr, qn('w:tc'))
            tcPr0 = etree.SubElement(tc0, qn('w:tcPr'))
            tcW0 = etree.SubElement(tcPr0, qn('w:tcW'))
            tcW0.set(qn('w:w'), str(CHART_TABLE_COL0_WIDTH))
            tcW0.set(qn('w:type'), 'dxa')
            # Vertical center alignment for the chart
            vAlign0 = etree.SubElement(tcPr0, qn('w:vAlign'))
            vAlign0.set(qn('w:val'), 'center')

            # Paragraph containing the image — centered, no extra spacing
            img_para = etree.SubElement(tc0, qn('w:p'))
            img_pPr = etree.SubElement(img_para, qn('w:pPr'))
            img_spacing = etree.SubElement(img_pPr, qn('w:spacing'))
            img_spacing.set(qn('w:before'), '0')
            img_spacing.set(qn('w:after'), '0')
            img_spacing.set(qn('w:line'), '240')
            img_spacing.set(qn('w:lineRule'), 'auto')
            img_jc = etree.SubElement(img_pPr, qn('w:jc'))
            img_jc.set(qn('w:val'), 'center')

            # Add chart image using python-docx InlineShape
            chart_path = chart_paths.get(chart_name)
            if chart_path and os.path.exists(chart_path):
                img_cx = CHART_IMG_CX
                img_cy = CHART_IMG_CY.get(chart_name, 3028869)

                # Add image to document part and get relationship ID
                # get_or_add_image returns (rId, Image) tuple
                rId, image_obj = doc.part.get_or_add_image(chart_path)

                # Build inline drawing XML
                drawing = _build_inline_image_xml(
                    rId, img_cx, img_cy,
                    f'{chart_name.lower()}_weekly_chart',
                    f'{chart_name} Weekly Candlestick Chart'
                )

                # Create a run and append drawing
                img_run = etree.SubElement(img_para, qn('w:r'))
                img_run.append(drawing)

                print(f"  [OK] {chart_name} chart image inserted ({img_cx}×{img_cy} EMU)")
            else:
                # Fallback — empty cell with placeholder text
                img_run = etree.SubElement(img_para, qn('w:r'))
                img_t = etree.SubElement(img_run, qn('w:t'))
                img_t.text = f"[{chart_name} Chart]"
                print(f"  [WARN] {chart_name} chart image not found at {chart_path}")

            # ─── Cell 1: Commentary ───
            tc1 = etree.SubElement(tr, qn('w:tc'))
            tcPr1 = etree.SubElement(tc1, qn('w:tcPr'))
            tcW1 = etree.SubElement(tcPr1, qn('w:tcW'))
            tcW1.set(qn('w:w'), str(CHART_TABLE_COL1_WIDTH))
            tcW1.set(qn('w:type'), 'dxa')

            # Get commentary
            outlook_data = tech_outlook.get(chart_name, {})
            commentary = outlook_data.get('commentary', [])

            # Heading paragraph: "{INDEX} (WEEKLY)"
            heading_para = etree.SubElement(tc1, qn('w:p'))
            h_pPr = etree.SubElement(heading_para, qn('w:pPr'))
            h_pStyle = etree.SubElement(h_pPr, qn('w:pStyle'))
            h_pStyle.set(qn('w:val'), 'MHeading2')
            h_spacing = etree.SubElement(h_pPr, qn('w:spacing'))
            h_spacing.set(qn('w:before'), '0')

            h_run = etree.SubElement(heading_para, qn('w:r'))
            h_rPr = etree.SubElement(h_run, qn('w:rPr'))
            h_sz = etree.SubElement(h_rPr, qn('w:sz'))
            h_sz.set(qn('w:val'), '22')  # 11pt — readable heading
            h_szCs = etree.SubElement(h_rPr, qn('w:szCs'))
            h_szCs.set(qn('w:val'), '22')
            h_color = etree.SubElement(h_rPr, qn('w:color'))
            h_color.set(qn('w:val'), HEADING_COLOR)
            h_b = etree.SubElement(h_rPr, qn('w:b'))
            h_bCs = etree.SubElement(h_rPr, qn('w:bCs'))
            h_t = etree.SubElement(h_run, qn('w:t'))
            h_t.text = display_headers.get(chart_name, chart_name)

            # Bullet paragraphs (commentary)
            for bullet_idx, bullet_text in enumerate(commentary):
                bp = etree.SubElement(tc1, qn('w:p'))
                bp_pPr = etree.SubElement(bp, qn('w:pPr'))

                is_last = (bullet_idx == len(commentary) - 1)

                # All commentary should use MBullet style for uniform formatting
                bp_style = etree.SubElement(bp_pPr, qn('w:pStyle'))
                bp_style.set(qn('w:val'), 'MBullet')

                # Run with 9.5pt font — readable in commentary column
                bp_run = etree.SubElement(bp, qn('w:r'))
                bp_rPr = etree.SubElement(bp_run, qn('w:rPr'))
                bp_sz = etree.SubElement(bp_rPr, qn('w:sz'))
                bp_sz.set(qn('w:val'), '19')   # 9.5pt
                bp_szCs = etree.SubElement(bp_rPr, qn('w:szCs'))
                bp_szCs.set(qn('w:val'), '19')
                bp_color = etree.SubElement(bp_rPr, qn('w:color'))
                bp_color.set(qn('w:val'), '000000')

                # Last bullet (S/R levels) is bold
                if is_last:
                    bp_b = etree.SubElement(bp_rPr, qn('w:b'))
                    bp_bCs = etree.SubElement(bp_rPr, qn('w:bCs'))

                bp_t = etree.SubElement(bp_run, qn('w:t'))
                bp_t.text = bullet_text
                bp_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

        print(f"  [OK] Chart + Commentary table: {len(chart_order)} rows created")
    else:
        print("  [SKIP] No tech outlook data — Page 2 chart table not generated")

    # ─── C) Page break → Page 3 starts here ───
    # Insert an explicit page break paragraph so ERT/Disclaimer goes to Page 3
    pg3_break_para = etree.Element(qn('w:p'))
    _insert_before_sectpr(pg3_break_para)
    pg3_break_pPr = etree.SubElement(pg3_break_para, qn('w:pPr'))
    pg3_spacing = etree.SubElement(pg3_break_pPr, qn('w:spacing'))
    pg3_spacing.set(qn('w:before'), '0')
    pg3_spacing.set(qn('w:after'), '0')
    pg3_break_run = etree.SubElement(pg3_break_para, qn('w:r'))
    pg3_break_br = etree.SubElement(pg3_break_run, qn('w:br'))
    pg3_break_br.set(qn('w:type'), 'page')
    print("  [OK] Page break inserted before Page 3 (contacts/disclaimer)")

    # ─── C) "EQUITY RESEARCH TEAM" heading ───
    print("\n[6/6] Adding footer elements (Page 3)...")

    ert_para = etree.Element(qn('w:p'))
    _insert_before_sectpr(ert_para)

    ert_pPr = etree.SubElement(ert_para, qn('w:pPr'))
    ert_spacing = etree.SubElement(ert_pPr, qn('w:spacing'))
    ert_spacing.set(qn('w:after'), '120')

    ert_run = etree.SubElement(ert_para, qn('w:r'))
    ert_rPr = etree.SubElement(ert_run, qn('w:rPr'))
    ert_sz = etree.SubElement(ert_rPr, qn('w:sz'))
    ert_sz.set(qn('w:val'), str(ERT_FONT_SIZE_HALFPTS))
    ert_szCs = etree.SubElement(ert_rPr, qn('w:szCs'))
    ert_szCs.set(qn('w:val'), str(ERT_FONT_SIZE_HALFPTS))
    ert_b = etree.SubElement(ert_rPr, qn('w:b'))
    ert_bCs = etree.SubElement(ert_rPr, qn('w:bCs'))
    ert_color = etree.SubElement(ert_rPr, qn('w:color'))
    ert_color.set(qn('w:val'), ERT_COLOR)
    ert_t = etree.SubElement(ert_run, qn('w:t'))
    ert_t.text = "EQUITY RESEARCH TEAM"

    print("  [OK] 'EQUITY RESEARCH TEAM' heading inserted")

    # ─── D) Contact table ───
    contact_tbl = etree.Element(qn('w:tbl'))
    _insert_before_sectpr(contact_tbl)

    ct_tblPr = etree.SubElement(contact_tbl, qn('w:tblPr'))
    ct_tblW = etree.SubElement(ct_tblPr, qn('w:tblW'))
    ct_tblW.set(qn('w:w'), '5000')
    ct_tblW.set(qn('w:type'), 'pct')

    # Grid columns
    ct_grid = etree.SubElement(contact_tbl, qn('w:tblGrid'))
    for gc_w in CONTACT_TABLE_GRID_COLS:
        gc = etree.SubElement(ct_grid, qn('w:gridCol'))
        gc.set(qn('w:w'), str(gc_w))

    # Header row
    ct_headers = ['Name', 'Designation', 'Email', 'Landline No.']
    ct_hdr_row = etree.SubElement(contact_tbl, qn('w:tr'))

    for col_idx, hdr_text in enumerate(ct_headers):
        tc = etree.SubElement(ct_hdr_row, qn('w:tc'))
        tcPr = etree.SubElement(tc, qn('w:tcPr'))
        tcW = etree.SubElement(tcPr, qn('w:tcW'))
        tcW.set(qn('w:w'), str(CONTACT_TABLE_CELL_WIDTHS[col_idx]))
        tcW.set(qn('w:type'), 'dxa')

        p = etree.SubElement(tc, qn('w:p'))
        r = etree.SubElement(p, qn('w:r'))
        rPr = etree.SubElement(r, qn('w:rPr'))
        sz_el = etree.SubElement(rPr, qn('w:sz'))
        sz_el.set(qn('w:val'), str(CONTACT_TABLE_FONT_SIZE))
        szCs_el = etree.SubElement(rPr, qn('w:szCs'))
        szCs_el.set(qn('w:val'), str(CONTACT_TABLE_FONT_SIZE))
        t = etree.SubElement(r, qn('w:t'))
        t.text = hdr_text

    # Data rows
    for person in RESEARCH_TEAM:
        ct_data_row = etree.SubElement(contact_tbl, qn('w:tr'))
        trPr = etree.SubElement(ct_data_row, qn('w:trPr'))
        trH = etree.SubElement(trPr, qn('w:trHeight'))
        trH.set(qn('w:val'), str(CONTACT_TABLE_ROW1_HEIGHT))

        values = [person['name'], person['designation'],
                  person['email'], person['phone']]

        for col_idx, val in enumerate(values):
            tc = etree.SubElement(ct_data_row, qn('w:tc'))
            tcPr = etree.SubElement(tc, qn('w:tcPr'))
            tcW = etree.SubElement(tcPr, qn('w:tcW'))
            tcW.set(qn('w:w'), str(CONTACT_TABLE_CELL_WIDTHS[col_idx]))
            tcW.set(qn('w:type'), 'dxa')

            p = etree.SubElement(tc, qn('w:p'))
            r = etree.SubElement(p, qn('w:r'))
            rPr = etree.SubElement(r, qn('w:rPr'))
            sz_el = etree.SubElement(rPr, qn('w:sz'))
            sz_el.set(qn('w:val'), str(CONTACT_TABLE_FONT_SIZE))
            szCs_el = etree.SubElement(rPr, qn('w:szCs'))
            szCs_el.set(qn('w:val'), str(CONTACT_TABLE_FONT_SIZE))
            # Name column is bold
            if col_idx == 0:
                b_el = etree.SubElement(rPr, qn('w:b'))
            t = etree.SubElement(r, qn('w:t'))
            t.text = val
            t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

    print("  [OK] Contact table inserted")

    # ─── E) Empty paragraph ───
    empty_para = etree.Element(qn('w:p'))
    _insert_before_sectpr(empty_para)

    # ─── F) Disclaimer paragraph ───
    disc_para = etree.Element(qn('w:p'))
    _insert_before_sectpr(disc_para)

    disc_run = etree.SubElement(disc_para, qn('w:r'))
    disc_rPr = etree.SubElement(disc_run, qn('w:rPr'))
    disc_b = etree.SubElement(disc_rPr, qn('w:b'))
    disc_bCs = etree.SubElement(disc_rPr, qn('w:bCs'))
    disc_t = etree.SubElement(disc_run, qn('w:t'))
    disc_t.text = DISCLAIMER_TEXT

    # ─── G) Company text paragraph ───
    company_para = etree.Element(qn('w:p'))
    _insert_before_sectpr(company_para)

    comp_run = etree.SubElement(company_para, qn('w:r'))
    comp_rPr = etree.SubElement(comp_run, qn('w:rPr'))
    comp_sz = etree.SubElement(comp_rPr, qn('w:sz'))
    comp_sz.set(qn('w:val'), str(COMPANY_TEXT_FONT_SIZE))
    comp_szCs = etree.SubElement(comp_rPr, qn('w:szCs'))
    comp_szCs.set(qn('w:val'), str(COMPANY_TEXT_FONT_SIZE))
    comp_t = etree.SubElement(comp_run, qn('w:t'))
    comp_t.text = COMPANY_TEXT
    comp_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

    # ─── H) Final empty paragraph ───
    final_empty = etree.Element(qn('w:p'))
    _insert_before_sectpr(final_empty)

    print("  [OK] Disclaimer and company info inserted")

    # Save python-docx content
    doc.save(output_path)

    # ── Step 1 (deferred): Patch header date via zip AFTER doc.save() ──
    # python-docx rewrites all package files on save, which would destroy
    # the gradient header design if we patched headers before saving.
    # Patching here preserves the full header design from the template.
    print("\n[1/6] Patching header date (post-save)...")
    month_pat = (r"(?:January|February|March|April|May|June|July|August|"
                 r"September|October|November|December)\s+\d{1,2},\s+\d{4}")

    with zipfile.ZipFile(output_path, 'r') as zin:
        files_out = {name: zin.read(name) for name in zin.namelist()}

    # Also pull the template's header files (which have the full gradient design)
    # and use those as the base, since python-docx may have simplified them
    with zipfile.ZipFile(TEMPLATE_PATH, 'r') as ztmpl:
        tmpl_files = {name: ztmpl.read(name) for name in ztmpl.namelist()}

    for fname in list(files_out.keys()):
        if re.match(r'word/header\d+\.xml', fname) or \
           re.match(r'word/_rels/header\d+\.xml\.rels', fname):
            # Use template's header (preserves gradient images) as base
            base = tmpl_files.get(fname, files_out[fname])
            txt = base.decode('utf-8')
            patched = re.sub(month_pat, header_date_str, txt)
            if patched != txt:
                print(f"  [OK] Date patched in {fname}")
            files_out[fname] = patched.encode('utf-8')

    # Also restore any media referenced by headers (e.g. gradient image, logo)
    for fname in tmpl_files:
        if fname.startswith('word/media/') and fname not in files_out:
            files_out[fname] = tmpl_files[fname]
            print(f"  [OK] Restored header media: {fname}")

    # ── CRITICAL FIX: Restore ALL sectPr header/footer references ──
    # python-docx rewrites section properties on save. We restore them from
    # the template: Section 1 (Page 1 gradient) and Section 2 (Page 2 plain).
    try:
        tmpl_doc_xml = tmpl_files.get('word/document.xml', b'').decode('utf-8', errors='replace')
        out_doc_xml = files_out.get('word/document.xml', b'').decode('utf-8', errors='replace')

        # Pull ALL sectPr blocks from the template
        tmpl_sectprs = list(re.finditer(
            r'<w:sectPr\b[^>]*>.*?</w:sectPr>', tmpl_doc_xml, re.DOTALL))

        # Page 1 sectPr = first one in template (has gradient header3, footer2/3)
        tmpl_sect1 = tmpl_sectprs[0].group(0) if len(tmpl_sectprs) >= 1 else None
        # Page 2 sectPr = last (final) one in template (has header4/5, footer4)
        tmpl_sect2 = tmpl_sectprs[-1].group(0) if len(tmpl_sectprs) >= 2 else None

        def _extract_refs(sect_xml):
            """Extract all headerReference and footerReference tags."""
            hrefs = re.findall(r'<w:headerReference[^>]*/>', sect_xml)
            frefs = re.findall(r'<w:footerReference[^>]*/>', sect_xml)
            title_pg = '<w:titlePg/>' in sect_xml or '<w:titlePg />' in sect_xml
            return hrefs, frefs, title_pg

        def _inject_refs(sect_xml, hrefs, frefs, title_pg):
            """Remove old refs and inject new ones; ensure titlePg if needed."""
            # Remove old header/footer refs and titlePg
            sect_xml = re.sub(r'<w:headerReference[^>]*/>\s*', '', sect_xml)
            sect_xml = re.sub(r'<w:footerReference[^>]*/>\s*', '', sect_xml)
            sect_xml = re.sub(r'<w:titlePg\s*/>\s*', '', sect_xml)
            # Build injection block
            inject = ''.join(hrefs) + ''.join(frefs)
            if title_pg:
                inject += '<w:titlePg/>'
            # Insert right after the opening <w:sectPr ...> tag
            insert_at = sect_xml.index('>') + 1
            return sect_xml[:insert_at] + inject + sect_xml[insert_at:]

        # Find all sectPr blocks in the OUTPUT document
        out_sectprs = list(re.finditer(
            r'<w:sectPr\b[^>]*>.*?</w:sectPr>', out_doc_xml, re.DOTALL))

        print(f"  [INFO] Found {len(out_sectprs)} sectPr block(s) in output document, "
              f"{len(tmpl_sectprs)} in template")

        if tmpl_sect1 and len(out_sectprs) >= 1:
            # Fix Section 1 (Page 1): use gradient header refs from template
            hrefs1, frefs1, tp1 = _extract_refs(tmpl_sect1)
            new_sect1 = _inject_refs(out_sectprs[0].group(0), hrefs1, frefs1, tp1)
            out_doc_xml = out_doc_xml[:out_sectprs[0].start()] + new_sect1 + out_doc_xml[out_sectprs[0].end():]
            print(f"  [OK] Restored Page 1 sectPr: gradient header3 now active for first page")

        # Re-find sectPrs after modification
        out_sectprs2 = list(re.finditer(
            r'<w:sectPr\b[^>]*>.*?</w:sectPr>', out_doc_xml, re.DOTALL))

        if tmpl_sect2 and len(out_sectprs2) >= 2:
            # Fix Section 2 (Page 2+): use plain header4/5 and footer4 from template
            hrefs2, frefs2, tp2 = _extract_refs(tmpl_sect2)
            # Add default footer ref if missing (so page 2 also shows footer)
            has_default_footer = any('type="default"' in f or 'w:type="default"' in f for f in frefs2)
            if not has_default_footer and frefs2:
                # Reuse the first footer as default too
                default_fr = frefs2[0].replace('w:type="first"', 'w:type="default"')
                frefs2.append(default_fr)
            new_sect2 = _inject_refs(out_sectprs2[-1].group(0), hrefs2, frefs2, tp2)
            out_doc_xml = out_doc_xml[:out_sectprs2[-1].start()] + new_sect2 + out_doc_xml[out_sectprs2[-1].end():]
            print(f"  [OK] Restored Page 2 sectPr: plain header4/5 + footer4 active for page 2+")
        elif len(out_sectprs2) < 2:
            print(f"  [WARN] Only {len(out_sectprs2)} sectPr found in output — Page 2 header refs NOT restored.")
            print(f"         This means the pb_para sectPr may not have survived doc.save().")

        files_out['word/document.xml'] = out_doc_xml.encode('utf-8')

    except Exception as _e:
        import traceback
        print(f"  [WARN] Could not restore sectPr header refs: {_e}")
        traceback.print_exc()

    # ── Restore ALL footer files from template ──
    for fname in tmpl_files:
        if re.match(r'word/footer\d+\.xml', fname):
            files_out[fname] = tmpl_files[fname]
        if re.match(r'word/_rels/footer\d+\.xml\.rels', fname):
            files_out[fname] = tmpl_files[fname]
        if fname == '[Content_Types].xml' and fname not in files_out:
            files_out[fname] = tmpl_files[fname]

        # Always use template's Content_Types.xml so all header/footer parts are registered
        files_out['[Content_Types].xml'] = tmpl_files['[Content_Types].xml']
        # NOTE: do NOT replace document.xml.rels here — the output version already
        # has all header/footer rIds, plus the chart image rIds added by python-docx.
        # Replacing it with the template version would erase the chart image entries.

    with zipfile.ZipFile(output_path, 'w', compression=zipfile.ZIP_DEFLATED) as zout:
        for fname, data in files_out.items():
            zout.writestr(fname, data)

    # Validate
    try:
        Document(output_path)
        print("\n  [OK] Document structure validated")
    except Exception as e:
        print(f"\n  [ERROR] Generated document may be corrupted: {e}")
        raise

# ══════════════════════════════════════════════════════════════════════════════
# MANUAL DERIVATIVES OVERRIDE — VALIDATION HELPERS
# ══════════════════════════════════════════════════════════════════════════════
#
# The automated paths (parse_bhavcopy_derivatives / fetch_option_chain_live)
# both compute PCR as sum(Put OI) / sum(Call OI) across ALL expiries — the
# firm's "cumulative PCR" definition. If both fail and we drop to manual entry,
# these helpers add guardrails without blocking entry outright: they warn and
# ask for an explicit confirmation before accepting a suspicious value.

_STRIKE_STEP = {"NIFTY": 50, "BANK NIFTY": 100}
_PCR_SANITY_LO, _PCR_SANITY_HI = 0.3, 2.5


def _get_spot_for_index(idx_name: str, mkt_data: Dict[str, Any]) -> Any:
    """Look up the current close for NIFTY / BANK NIFTY from already-fetched
    mkt_data['ema'] for use as a soft sanity check. Returns None if absent."""
    ema_name_map = {"NIFTY": "NIFTY", "BANK NIFTY": "BANK NIFTY"}
    target = ema_name_map.get(idx_name)
    for row in mkt_data.get("ema", []):
        if row.get("name") == target:
            return row.get("close")
    return None


def _validate_manual_strike(idx_name: str, label: str, raw: str,
                             spot: Any = None) -> bool:
    """Validate a manually typed strike.

    Checks: (1) parses as int, (2) is a multiple of the index's strike step,
    (3) isn't wildly far from spot.  Warns and asks for YES-to-override on any
    failure.  Returns True if the value is acceptable.
    """
    if not raw:
        return True  # empty means "skip this field" — caller handles it

    try:
        strike = int(raw.replace(",", ""))
    except ValueError:
        print(f"    [REJECTED] '{raw}' is not a valid integer strike.")
        return False

    step = _STRIKE_STEP.get(idx_name)
    if step and strike % step != 0:
        print(f"    [WARN] {label} = {strike:,} is NOT a multiple of {step}. "
              f"{idx_name} strikes must be — this is usually a call/put side "
              f"swap or a value from the wrong index.")
        if input("    Type YES to keep it anyway, or press ENTER to re-enter: "
                 ).strip().upper() != "YES":
            return False

    if spot:
        pct = abs(strike - spot) / spot * 100
        if pct > 15:
            print(f"    [WARN] {label} = {strike:,} is {pct:.1f}% away from "
                  f"current {idx_name} close (~{spot:,.0f}). Check that this "
                  f"isn't from the wrong index, expiry, or date.")
            if input("    Type YES to keep it anyway, or press ENTER to re-enter: "
                     ).strip().upper() != "YES":
                return False

    return True


def _validate_manual_pcr(idx_name: str, raw: str) -> Any:
    """Validate a manually typed PCR value.

    Returns the float if valid (and confirmed if suspicious), or None to
    signal that the caller should re-prompt.
    """
    try:
        pcr = float(raw)
    except ValueError:
        print(f"    [REJECTED] '{raw}' is not a valid PCR number.")
        return None

    if not (_PCR_SANITY_LO <= pcr <= _PCR_SANITY_HI):
        print(f"    [WARN] PCR {pcr:.2f} for {idx_name} is outside the expected "
              f"{_PCR_SANITY_LO}–{_PCR_SANITY_HI} range.")
        print(f"    REMINDER: this is CUMULATIVE PCR across ALL expiries "
              f"(total Put OI ÷ total Call OI). Near-expiry-only PCR from "
              f"trading terminals reads differently and is the most common "
              f"source of a wrong number here.")
        if input("    Type YES to keep it anyway, or press ENTER to re-enter: "
                 ).strip().upper() != "YES":
            return None

    return pcr


def main():
    """Main entry point."""
    print("\n" + "=" * 70)
    print(" WEEKLY EQUITY REPORT GENERATOR - PRODUCTION GRADE")
    print(" Monarch Networth Capital")
    print("=" * 70)

    # Check template exists
    if not os.path.exists(TEMPLATE_PATH):
        print(f"\n[ERROR] Template not found at: {TEMPLATE_PATH}")
        sys.exit(1)

    # Get date range — only ask for the END Sunday; start is always 7 days before
    print("\nWeekly report covers exactly one week (Sunday to Sunday).")
    print("Enter the LAST Sunday of the report week (DD-MM-YYYY):\n")

    while True:
        end_str = input("END Sunday (DD-MM-YYYY): ").strip()

        try:
            start_date, end_date = validate_end_sunday(end_str)
            print(f"\n[OK] Week: {start_date.strftime('%A, %d %B %Y')} → "
                  f"{end_date.strftime('%A, %d %B %Y')}")
            break
        except ValueError as e:
            print(f"\n[ERROR] {e}")
            if input("\nTry again? (y/n): ").strip().lower() != 'y':
                sys.exit(0)
            print()

    # Fetch all data
    try:
        mkt_data = fetch_all_market_data(start_date, end_date)
    except Exception as e:
        print(f"\n[FATAL] Market data fetch failed: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)

    # FII/DII
    try:
        prev_friday, curr_friday = get_friday_dates(start_date, end_date)
        
        expected_days = 5
        nifty_df = mkt_data.get("yf_cache", {}).get("NIFTY 50")
        if nifty_df is not None and not nifty_df.empty:
            import pandas as pd
            mask = (nifty_df.index > pd.Timestamp(prev_friday)) & (nifty_df.index <= pd.Timestamp(curr_friday))
            actual_days = len(nifty_df[mask])
            if actual_days > 0:
                expected_days = actual_days

        print("\n" + "-" * 70)
        print(f"FII/DII for week ending {curr_friday.strftime('%d-%b-%Y')} (automated)")
        print("-" * 70)
        
        single_day = get_fii_dii_for_date(curr_friday)
        weekly = get_weekly_fii_dii(prev_friday, curr_friday, expected_days)

        if weekly.get("is_complete"):
            fii_dii = weekly
        elif single_day.get("fii") is not None:
            days = weekly.get("days_covered", 0)
            fii_dii = {
                "fii": single_day["fii"],
                "dii": single_day["dii"],
                "is_weekly": weekly.get("is_weekly", False),
                "days_covered": max(days, 1),
                "is_complete": False,
            }
            if single_day.get("source") == "log":
                print(f"  [INFO] Using previously logged value for {curr_friday.strftime('%d-%b-%Y')}.")
        else:
            print(f"  [WARN] No FII/DII figure available for {curr_friday.strftime('%d-%b-%Y')} "
                  f"— it's neither today nor a previously logged date. Run the generator "
                  f"on or near that Friday once to capture it automatically going forward.")
            fii_dii = {"fii": None, "dii": None, "is_weekly": False,
                       "days_covered": 0, "is_complete": False}

        if fii_dii.get("is_weekly") and not fii_dii.get("is_complete") and fii_dii.get("fii") is not None:
            expected = fii_dii.get("expected_days", 5)
            print(f"  [INFO] Partial week ({fii_dii['days_covered']}/{expected} days logged). "
                  f"Total completes automatically as remaining trading days are logged.")

    except Exception as e:
        print(f"\n[WARN] FII/DII automation failed: {e}")
        fii_dii = {"fii": None, "dii": None, "is_weekly": False,
                   "days_covered": 0, "is_complete": False}

    # Manual Fallback if automated data is missing or incomplete
    if not fii_dii.get("is_weekly") or not fii_dii.get("is_complete"):
        print("\n" + "-" * 70)
        print("FII/DII MANUAL OVERRIDE (Automated weekly data incomplete)")
        print("-" * 70)
        print("You can enter the correct WEEKLY totals manually.")
        print("(Press ENTER to skip and use whatever data was found automatically)\n")
        
        try:
            fii_input = input(f"  Weekly FII net value in Cr [current: {fii_dii.get('fii')}]: ").strip()
            dii_input = input(f"  Weekly DII net value in Cr [current: {fii_dii.get('dii')}]: ").strip()
            
            if fii_input:
                fii_dii["fii"] = int(float(fii_input.replace(",", "")))
                fii_dii["is_weekly"] = True
                fii_dii["is_complete"] = True
                print(f"  [OK] FII overridden to: {fii_dii['fii']:,} Cr")
            
            if dii_input:
                fii_dii["dii"] = int(float(dii_input.replace(",", "")))
                fii_dii["is_weekly"] = True
                fii_dii["is_complete"] = True
                print(f"  [OK] DII overridden to: {fii_dii['dii']:,} Cr")
        except Exception as e:
            print(f"  [WARN] Override input error: {e} — using fetched value.")


    # Support/Resistance manual override
    # Only manual input is accepted for S1/S2/R1/R2. If skipped, they are left blank.
    # Close value is kept automatically.
    print("\n" + "-" * 70)
    print("SUPPORT / RESISTANCE MANUAL OVERRIDE")
    print("-" * 70)
    print("Auto-calculated S/R levels are shown in brackets.")
    print("Press ENTER to use the auto-calculated value, or type a new value.\n")

    def _sr_input(label, default):
        """Prompt user for a single S/R level; returns user value or None if left blank."""
        default_str = f"{default:,}" if default is not None else "blank"
        try:
            val = input(f"    {label} [auto: {default_str}]: ").strip()
            if val == "":
                return None
            if val.lower() == "auto":
                return default
            return int(float(val.replace(",", "")))
        except Exception:
            return None

    for sr_row in mkt_data["sr"]:
        name = sr_row["name"]
        print(f"  --- {name} (Close: {sr_row.get('close')}) ---")
        sr_row["s2"] = _sr_input("S2", sr_row.get("s2"))
        sr_row["s1"] = _sr_input("S1", sr_row.get("s1"))
        sr_row["r1"] = _sr_input("R1", sr_row.get("r1"))
        sr_row["r2"] = _sr_input("R2", sr_row.get("r2"))
        
        # If no manual input for S/R, blank the bias column
        if sr_row["s2"] is None and sr_row["s1"] is None and sr_row["r1"] is None and sr_row["r2"] is None:
            sr_row["bias"] = ""
            
        # Close remains untouched
        print(f"    => S2={sr_row['s2']}  S1={sr_row['s1']}  "
              f"Close={sr_row['close']}  R1={sr_row['r1']}  R2={sr_row['r2']}\n")

    # Constituents
    print("\n[INFO] Fetching constituent data...")
    try:
        prev_fri, curr_fri = mkt_data["actual_fridays"]
        constituents = fetch_constituents(prev_fri, curr_fri)
    except Exception as e:
        print(f"\n[WARN] Constituents fetch failed: {e}")
        constituents = None

    # Global markets
    print("\n[INFO] Fetching global markets...")
    try:
        prev_fri_tgt, curr_fri_tgt = mkt_data["target_fridays"]
        global_mkts = fetch_global_markets(prev_fri_tgt, curr_fri_tgt)
    except Exception as e:
        print(f"\n[WARN] Global markets fetch failed: {e}")
        global_mkts = None

    # Derivatives manual override / automation
    print("\n" + "-" * 70)
    print("DERIVATIVES DATA EXTRACTION")
    print("-" * 70)
    
    # 1. Try BhavCopy First
    _, curr_friday_actual = mkt_data["actual_fridays"]
    derivatives_data = parse_bhavcopy_derivatives(end_date, curr_friday_actual)
    
    # 2. Fallback to Live NSE Option Chain
    if derivatives_data:
        print("  [SUCCESS] Derivatives data successfully automated from Bhavcopy!")
    else:
        print("  [INFO] Bhavcopy parse failed or no data. Falling back to Live Option Chain...")
        derivatives_data = fetch_option_chain_live(end_date, curr_friday_actual)
        
        if derivatives_data:
            print("  [SUCCESS] Derivatives data successfully automated from Live NSE Option Chain!")
        else:
            print("\n  [WARN] Could not automate derivatives. Falling back to MANUAL OVERRIDE.")
            print("  Enter the Max Call OI, Max Put OI, and PCR values.")
            print("  Press ENTER to skip and leave derivatives blank in the report.\n")
            
            derivatives_data = {}
            for idx_name in ["NIFTY", "BANK NIFTY"]:
                print(f"  --- {idx_name} ---")
                spot = _get_spot_for_index(idx_name, mkt_data)
                if spot:
                    print(f"    (Current {idx_name} close: {spot:,.0f} — "
                          f"strikes should be within ~15% of this)")
                try:
                    while True:
                        oi_add_c = input("    Max OI Addition Call (e.g. 25000): ").strip()
                        if _validate_manual_strike(idx_name, "Max OI Addition Call",
                                                   oi_add_c, spot):
                            break

                    while True:
                        oi_add_p = input("    Max OI Addition Put  (e.g. 24000): ").strip()
                        if _validate_manual_strike(idx_name, "Max OI Addition Put",
                                                   oi_add_p, spot):
                            break

                    while True:
                        c_oi = input("    Max Call OI Strike   (e.g. 24800): ").strip()
                        if _validate_manual_strike(idx_name, "Max Call OI Strike",
                                                   c_oi, spot):
                            break

                    while True:
                        p_oi = input("    Max Put  OI Strike   (e.g. 23000): ").strip()
                        if _validate_manual_strike(idx_name, "Max Put OI Strike",
                                                   p_oi, spot):
                            break

                    expiry = input("    Expiry Date String    (e.g. 21st Jul): ").strip()

                    pcr = None
                    while True:
                        pcr_in = input("    Cumulative PCR all-expiry (e.g. 1.25): ").strip()
                        if not pcr_in:
                            break  # empty = skip; handled by completeness check below
                        pcr = _validate_manual_pcr(idx_name, pcr_in)
                        if pcr is not None:
                            break

                    pcr_date = input("    PCR Date String       (e.g. Jul 18): ").strip()

                    if oi_add_c and oi_add_p and c_oi and p_oi and pcr is not None \
                            and expiry and pcr_date:
                        derivatives_data[idx_name] = {
                            "oi_add_c": oi_add_c.replace(",", ""),
                            "oi_add_p": oi_add_p.replace(",", ""),
                            "call_oi":  c_oi.replace(",", ""),
                            "put_oi":   p_oi.replace(",", ""),
                            "pcr":      pcr,
                            "expiry":   expiry,
                            "pcr_date": pcr_date,
                            "source":   "manual",  # traceable in logs if queried later
                        }
                        print(f"    => PCR={pcr:.2f} recorded (source: manual).\n")
                    else:
                        print(f"    => Skipped {idx_name} derivatives data.\n")
                except Exception as e:
                    print(f"    [WARN] Invalid input, skipping {idx_name}: {e}\n")

    # Build narrative
    narrative = build_narrative(mkt_data, fii_dii, constituents, global_mkts,
                                start_date, end_date, derivatives_data)

    # Generate output filename
    output_filename = (f"Weekly_Equity_Report_"
                      f"{start_date.strftime('%d%b%Y')}_to_"
                      f"{end_date.strftime('%d%b%Y')}.docx")
    output_path = os.path.join(os.path.dirname(TEMPLATE_PATH), output_filename)

    # Handle locked file
    if os.path.exists(output_path):
        try:
            with open(output_path, 'a'):
                pass
        except OSError:
            for i in range(1, 10):
                alt_path = output_path.replace('.docx', f'_{i}.docx')
                if not os.path.exists(alt_path):
                    output_path = alt_path
                    print(f"\n[INFO] Original file locked — using {os.path.basename(output_path)}")
                    break

    # ═══════════════════════════════════════════════════════════════
    # PAGE 2: CHARTS & TECHNICAL OUTLOOK (Automated)
    # ═══════════════════════════════════════════════════════════════
    print("\n" + "=" * 70)
    print("GENERATING PAGE 2 (TECHNICAL OUTLOOK & CHARTS)")
    print("=" * 70)
    
    try:
        import chart_generator
        
        yf_cache = mkt_data.get("yf_cache", {})
        
        # Build resolved_closes so chart generator uses the correct authoritative
        # close for indices like FINNIFTY where YF returns Close=NaN on the
        # current Friday.  These closes were already resolved by fetch_all_market_data.
        resolved_closes = {
            row["name"]: row["close"]
            for row in mkt_data.get("indices", [])
            if row.get("close") is not None
        }
        
        # 1. Generate chart images
        chart_paths = chart_generator.generate_all_charts(
            yf_cache, end_date, resolved_closes=resolved_closes
        )
        
        # 2. Generate technical commentary
        tech_outlook = chart_generator.generate_all_technical_data(
            yf_cache=yf_cache,
            sr_rows=mkt_data["sr"],
            indices_data=mkt_data["indices"],
            end_date=end_date,
            resolved_closes=resolved_closes,
        )
        print("  [OK] Page 2 data generated successfully.")
    except Exception as e:
        print(f"  [ERROR] Page 2 generation failed: {e}")
        import traceback
        traceback.print_exc()
        chart_paths = None
        tech_outlook = None

    # Fill document
    try:
        fill_docx_document(
            mkt_data=mkt_data, 
            fii_dii=fii_dii, 
            narrative=narrative, 
            start_date=start_date, 
            end_date=end_date, 
            output_path=output_path,
            tech_outlook=tech_outlook,
            chart_paths=chart_paths
        )
    except Exception as e:
        print(f"\n[FATAL] Document generation failed: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)

    # Data quality summary
    dq = mkt_data["data_quality"]
    print("\n" + "=" * 70)
    print("DATA QUALITY SUMMARY")
    print("=" * 70)
    print(f"  Indices: {dq['indices_ok']}/{dq['indices_total']} complete")
    print(f"  Sectors: {dq['sectors_ok']}/{dq['sectors_total']} complete")
    if fii_dii.get("is_complete") and fii_dii.get("is_weekly"):
        expected = fii_dii.get("expected_days", 5)
        fii_status = f"✓ (Complete {expected}-day automated weekly sum)"
    elif fii_dii.get("is_weekly"):
        expected = fii_dii.get("expected_days", 5)
        fii_status = f"✓ (Partial automated weekly sum - {fii_dii.get('days_covered')}/{expected} days)"
    elif fii_dii.get('fii') is not None:
        fii_status = "✓ (Single latest trading day snapshot)"
    else:
        fii_status = "✗ unavailable"
        
    print(f"  FII/DII: {fii_status}")
    print(f"  Global:  {'✓' if global_mkts else '✗'}")

    # Success
    print("\n" + "=" * 70)
    print(f"  [SUCCESS] Report generated: {os.path.basename(output_path)}")
    print("=" * 70)


if __name__ == "__main__":
    main()